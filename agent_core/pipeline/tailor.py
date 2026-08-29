"""Tailor resume for a specific job using LLM. Output .docx + .md + open job link.

Anti-hallucination strategy (B+A):
- Hard facts (dates, numbers) are extracted from the original resume, injected
  into the prompt as must-preserve items (A), and verified in the output (B).
- Low temperature (0.1) to reduce creative drift.
- See extract_hard_facts / verify_facts / diff_resumes.
"""

import difflib
import logging
import re
import webbrowser
from pathlib import Path

from agent_core.config import load_resume
from agent_core.llm.providers import call_llm_with_retry
from agent_core.pipeline.text_utils import has_all_sections

logger = logging.getLogger(__name__)

# JD over this many chars is truncated (with a marker) to keep the prompt bounded.
MAX_JD_CHARS = 8000

# Low temperature: tailoring is rephrase+reorder, not invention.
TAILOR_TEMPERATURE = 0.1

# Hard facts that must be preserved verbatim.
# Dates: 2020.07 / 2020-07 / 2020年7月 / 2020 / 至今 / present
DATE_PATTERN = re.compile(
    r"\d{4}\s*[.\-/年]\s*\d{1,2}\s*月?|\d{4}\s*年|\b(?:19|20)\d{2}\b|至今|present",
    re.IGNORECASE,
)
# Numbers with units (%, 万, 年, 条, ...) and standalone 2+ digit numbers.
NUMBER_PATTERN = re.compile(r"\d+(?:\.\d+)?\s*[%％万千亿条个万人秒分小时年元]|\b\d{2,}\b")

TAILOR_PROMPT = (
    "You are a resume tailoring expert. Adapt the resume below for a specific job.\n"
    "\n"
    "Job: {title} @ {company} | {location}\n"
    "JD: {description}\n"
    "\n"
    "Original Resume:\n"
    "{resume}\n"
    "\n"
    "Rules:\n"
    "1. Keep ALL facts exactly as-is (dates, companies, titles, skills)\n"
    "2. Reorder/rephrase bullets to emphasize skills relevant to this job\n"
    "3. Update self-summary to highlight match with this role\n"
    "4. Do NOT invent any experience or skills\n"
    "5. The following hard facts MUST appear VERBATIM in the output — do not reword, "
    "reformat, translate, or drop any of them:\n"
    "{hard_facts}\n"
    "6. Output tailored resume in clean Markdown with sections: "
    "教育背景, 核心能力 (3-5 most relevant skills for this job), "
    "工作经历/项目经历 (reordered by relevance), 技能, 自我评价\n"
    "\n"
    "Return ONLY the Markdown resume. No other text."
)


def extract_hard_facts(text):
    """Extract hard facts (dates, numbers) that must be preserved verbatim.

    Returns {"dates": [...], "numbers": [...]} with deduplicated sorted lists.
    """
    dates = sorted({m.group(0).strip() for m in DATE_PATTERN.finditer(text)})
    date_set = {d.lower() for d in dates}
    numbers = sorted(
        {
            m.group(0).strip()
            for m in NUMBER_PATTERN.finditer(text)
            if m.group(0).strip().lower() not in date_set
        }
    )
    return {"dates": dates, "numbers": numbers}


def verify_facts(tailored_text, facts):
    """Return list of (kind, value) for hard facts missing from tailored_text.

    A missing fact likely means the LLM altered or dropped it (potential hallucination).
    Date matching is case-insensitive; numbers are matched as substrings.
    """
    lowered = tailored_text.lower()
    missing = []
    for d in facts["dates"]:
        if d.lower() not in lowered:
            missing.append(("date", d))
    for num in facts["numbers"]:
        if num not in tailored_text:
            missing.append(("number", num))
    return missing


def diff_resumes(original, tailored, max_lines=200):
    """Return a unified diff between original and tailored resume (truncated)."""
    diff = difflib.unified_diff(
        original.splitlines(keepends=False),
        tailored.splitlines(keepends=False),
        fromfile="original",
        tofile="tailored",
        lineterm="",
    )
    lines = list(diff)[:max_lines]
    return "\n".join(lines)


def _format_hard_facts(facts, max_items=60):
    """Render hard facts as a compact list for prompt injection."""
    items = [f"日期: {d}" for d in facts["dates"]] + [f"数字: {n}" for n in facts["numbers"]]
    if len(items) > max_items:
        items = items[:max_items] + [f"...(+{len(items) - max_items} more)"]
    return "\n".join(items) if items else "(none)"


async def tailor_resume(
    job, config, llm_provider, direction=None, feedback=None, timeout: float | None = None
):
    """Generate a tailored resume for a job. Returns markdown text.

    Injects hard facts into the prompt (A) and verifies them in the output (B),
    logging a warning for any missing/changed facts. Optional `feedback` from a
    human reviewer is appended to steer regeneration.
    """
    if direction is None:
        direction = job.direction
    resume_text = load_resume(config, direction)
    facts = extract_hard_facts(resume_text)

    jd = job.description or ""
    jd_len = len(jd)
    if jd_len > MAX_JD_CHARS:
        # Explicit original length: don't rely on right-hand-side evaluation
        # order (len(jd) after the slice would report the truncated size).
        jd = jd[:MAX_JD_CHARS] + f"\n[...JD 已截断，原文共 {jd_len} 字符...]"
        logger.warning("JD truncated: %d -> %d chars", jd_len, MAX_JD_CHARS)

    prompt = TAILOR_PROMPT.format(
        title=job.title,
        company=job.company,
        location=job.location or config.search_location,
        description=jd,
        resume=resume_text,
        hard_facts=_format_hard_facts(facts),
    )
    if feedback:
        prompt += f"\n\n人工审核改进意见（请据此调整本次生成）：{feedback}"

    _REQUIRED_SECTIONS = ("## 教育背景", "## 核心能力", "## 工作经历", "## 技能", "## 自我评价")

    def _looks_truncated(text: str) -> bool:
        # Missing trailing sections is the strongest signal the LLM output was
        # cut off (common when thinking mode consumes the shared max_tokens).
        return not has_all_sections(text, _REQUIRED_SECTIONS)

    response = await call_llm_with_retry(
        llm_provider,
        messages=[{"role": "user", "content": prompt}],
        temperature=TAILOR_TEMPERATURE,
        max_tokens=config.llm.max_tokens,
        timeout=timeout,
    )
    tailored = response.strip()

    # Retry once with the SAME provider (thinking stays enabled as configured)
    # if the first pass looks truncated. This preserves output quality while
    # still giving the model a second chance to emit the full markdown.
    if _looks_truncated(tailored):
        logger.warning(
            "tailor_resume: first pass missing required sections (%d chars), retrying once",
            len(tailored),
        )
        try:
            retry = await call_llm_with_retry(
                llm_provider,
                messages=[{"role": "user", "content": prompt}],
                temperature=TAILOR_TEMPERATURE,
                max_tokens=config.llm.max_tokens,
                timeout=timeout,
            )
            retry_text = retry.strip()
            if len(retry_text) > len(tailored) and _looks_truncated(retry_text) is False:
                tailored = retry_text
            else:
                logger.warning("tailor_resume retry did not improve; keeping first pass")
        except Exception as e:  # noqa: BLE001
            logger.warning("tailor_resume retry failed: %s", e)

    missing = verify_facts(tailored, facts)
    if missing:
        logger.warning(
            "Hard-fact check: %d item(s) missing/changed: %s",
            len(missing),
            ", ".join(v for _, v in missing[:10]),
        )
    return tailored


def save_resume(text, job, output_dir="output"):
    """Save tailored resume as .md and .docx. Returns file paths."""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r'[\\/*?:"<>|]', "", job.title)[:30]
    safe_company = re.sub(r'[\\/*?:"<>|]', "", job.company)[:20]
    base = f"{output_dir}/{safe_company}_{safe_title}"

    md_path = f"{base}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(text)
    logger.info(f"Saved: {md_path}")

    docx_path = f"{base}.docx"
    _save_docx(text, docx_path)
    logger.info(f"Saved: {docx_path}")
    return {"md": md_path, "docx": docx_path}


def _split_table_row(line):
    """Split a markdown table row '| a | b |' into ['a', 'b']."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_inline_runs(paragraph, text):
    """Add text to a paragraph, rendering **bold**, *italic* and _italic_ inline."""
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*|_[^_]+?_)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _save_docx(text, path):
    """Convert markdown resume to .docx document.

    Supports: #/##/### headings, bullet/numbered lists, **bold**, *italic*,
    _italic_, GFM tables, and fenced code blocks.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    lines = text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].strip()

        # Fenced code block ```...```
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            continue

        # GFM table: a | header | row followed by a |---| separator row
        if (
            line.startswith("|")
            and line.endswith("|")
            and i + 1 < n
            and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip())
        ):
            header = _split_table_row(line)
            i += 2  # skip header + separator
            body_rows = []
            while i < n and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                body_rows.append(_split_table_row(lines[i].strip()))
                i += 1
            n_cols = max([len(header)] + [len(r) for r in body_rows])
            table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass  # template without that style; render borderless
            for j, h in enumerate(header):
                if j < n_cols:
                    table.rows[0].cells[j].text = h
            for r_idx, row in enumerate(body_rows, start=1):
                for j, c in enumerate(row):
                    if j < n_cols:
                        table.rows[r_idx].cells[j].text = c
            continue

        if not line:
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("- ") or line.startswith("* "):
            _add_inline_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        elif re.match(r"^\d+\.\s", line):
            _add_inline_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s*", "", line))
        else:
            _add_inline_runs(doc.add_paragraph(), line)
        i += 1

    doc.save(path)


def open_job_link(job):
    """Log every available platform URL and open the first one in the browser.

    Opens only the first URL to avoid spamming tabs; all platform URLs are logged
    so the user can open the others manually.
    """
    urls = job.urls if isinstance(job.urls, dict) else {}
    valid = [(platform, url) for platform, url in urls.items() if url]
    for platform, url in valid:
        logger.info("Job URL [%s]: %s", platform, url)
    if valid:
        logger.info("Opening: %s", valid[0][1])
        webbrowser.open(valid[0][1])
    return [url for _, url in valid]
